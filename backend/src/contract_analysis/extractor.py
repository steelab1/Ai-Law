"""
Contract Text Extractor
Trich xuat text tu file PDF va Word (.docx)
"""

import logging
import hashlib
from typing import Tuple
from io import BytesIO

logger = logging.getLogger(__name__)


class ContractExtractor:
    """Extract text from PDF and Word documents."""

    SUPPORTED_FORMATS = ['.pdf', '.docx', '.doc']
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

    @staticmethod
    def extract_from_pdf(file_bytes: bytes) -> str:
        """
        Extract text from PDF file using PyMuPDF.

        Args:
            file_bytes: Raw bytes of PDF file

        Returns:
            Extracted text with page markers
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF is required. Install with: pip install PyMuPDF")

        text_parts = []

        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page_num, page in enumerate(doc, 1):
                text = page.get_text("text")
                if text.strip():
                    text_parts.append(f"--- Trang {page_num} ---\n{text}")

        return "\n\n".join(text_parts)

    @staticmethod
    def extract_from_docx(file_bytes: bytes) -> str:
        """
        Extract text from Word document (.docx).

        Args:
            file_bytes: Raw bytes of Word file

        Returns:
            Extracted text from paragraphs and tables
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required. Install with: pip install python-docx")

        doc = Document(BytesIO(file_bytes))
        text_parts = []

        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    @classmethod
    def extract(cls, file_bytes: bytes, filename: str) -> Tuple[str, str]:
        """
        Extract text from uploaded file.

        Args:
            file_bytes: Raw bytes of the file
            filename: Original filename to determine type

        Returns:
            Tuple of (extracted_text, file_hash)

        Raises:
            ValueError: If file is too large, unsupported, or empty
        """
        # Validate file size
        if len(file_bytes) > cls.MAX_FILE_SIZE:
            raise ValueError(
                f"File qua lon. Kich thuoc toi da: {cls.MAX_FILE_SIZE // (1024 * 1024)}MB"
            )

        # Calculate hash for deduplication
        file_hash = hashlib.sha256(file_bytes).hexdigest()

        # Determine file type and extract
        filename_lower = filename.lower()

        if filename_lower.endswith('.pdf'):
            logger.info(f"Extracting text from PDF: {filename}")
            text = cls.extract_from_pdf(file_bytes)

        elif filename_lower.endswith(('.docx', '.doc')):
            logger.info(f"Extracting text from Word: {filename}")
            text = cls.extract_from_docx(file_bytes)

        else:
            raise ValueError(
                f"Dinh dang file khong duoc ho tro. "
                f"Cac dinh dang ho tro: {', '.join(cls.SUPPORTED_FORMATS)}"
            )

        # Validate extracted content
        if not text or len(text.strip()) < 100:
            raise ValueError(
                "Khong the trich xuat du noi dung tu file. "
                "Vui long kiem tra file khong bi hong hoac trong."
            )

        logger.info(f"Extracted {len(text)} characters from {filename}")
        return text, file_hash

    @staticmethod
    def get_file_info(file_bytes: bytes, filename: str) -> dict:
        """
        Get basic file information.

        Args:
            file_bytes: Raw bytes of the file
            filename: Original filename

        Returns:
            Dict with file metadata
        """
        return {
            "filename": filename,
            "size_bytes": len(file_bytes),
            "size_mb": round(len(file_bytes) / (1024 * 1024), 2),
            "hash": hashlib.sha256(file_bytes).hexdigest()[:16],
            "extension": filename.split('.')[-1].lower() if '.' in filename else "unknown"
        }
